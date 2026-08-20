from datetime import datetime
from docx import Document

# Build the conversation content
conversation = []
conversation.append(f"Document generated: {datetime.utcnow().isoformat()} UTC")
conversation.append("")
conversation.append("--- User-provided workspace info and requests ---")
conversation.append("")
conversation.append("Workspace folders and structure (truncated):")
conversation.append("- GitChat_Summary.md")
conversation.append("- PR_DRAFT_feature_payment-status.md")
conversation.append("- README.md")
conversation.append("- api/ (server.js, routes, data, etc.)")
conversation.append("- ui/ (src, public, tests, etc.)")
conversation.append("")
conversation.append("User request:")
conversation.append("make one word document of our important conversation and store the name as Satya_GitCopilot_Capstone_Project")
conversation.append("")
conversation.append("Context notes:")
conversation.append("- OS: Windows")
conversation.append("")
conversation.append("--- Conversation (selected important items) ---")
conversation.append("")
conversation.append("User: make one word document of our important conversation and store the name as Satya_GitCopilot_Capstone_Project")
conversation.append("")
conversation.append("Assistant: Creating this document and saving as Satya_GitCopilot_Capstone_Project.docx in the repository root.")

# Join into a single text block
body = "\n".join(conversation)

# Create the Word document
doc = Document()
doc.add_heading('Satya GitCopilot Capstone — Conversation Export', level=1)
doc.add_paragraph(body)

output_path = 'Satya_GitCopilot_Capstone_Project.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
